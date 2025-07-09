import json
import os
from typing import List, Dict, Any, Tuple
from dataclasses import dataclass
import re
from specs.disk import resolve_path
# Required libraries for docx generation
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. DOCX export will not be available.")
downloads = resolve_path()
nburl = os.path.join(downloads, "001-fundamentals.ipynb")

@dataclass
class CellData:
    """Data structure to hold cell information"""
    order: int
    content: str
    comment: str = ""

class NotebookParser:
    """Parser for Jupyter notebooks to extract and organize content"""
    
    def __init__(self, notebook_path: str):
        self.notebook_path = notebook_path
        self.cells_data: List[CellData] = []
    
    def parse_notebook(self, cell_type_filter: str = None) -> List[CellData]:
        """
        Parse the Jupyter notebook and extract content into a structured format
        
        Args:
            cell_type_filter: Optional filter for cell types ('code', 'markdown', 'raw')
                            If None, includes all cell types
        
        Returns:
            List[CellData]: List of cell data with order, content, and comments
        """
        if not os.path.exists(self.notebook_path):
            raise FileNotFoundError(f"Notebook file not found: {self.notebook_path}")
        
        try:
            with open(self.notebook_path, 'r', encoding='utf-8') as f:
                notebook = json.load(f)
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON in notebook file: {e}")
        
        cells = notebook.get('cells', [])
        self.cells_data = []
        
        for i, cell in enumerate(cells, 1):
            cell_type = cell.get('cell_type', '')
            source = cell.get('source', [])
            
            # Apply cell type filter if specified
            if cell_type_filter and cell_type != cell_type_filter:
                continue
            
            # Join source lines into a single string
            if isinstance(source, list):
                content = ''.join(source)
            else:
                content = str(source)
            
            # Extract comments based on cell type
            comment = self._extract_comment(cell, cell_type)
            
            # Only include cells with content
            if content.strip():
                cell_data = CellData(
                    order=i,
                    content=content.strip(),
                    comment=comment
                )
                self.cells_data.append(cell_data)
        
        return self.cells_data
    
    def _extract_comment(self, cell: Dict[Any, Any], cell_type: str) -> str:
        """
        Extract comments from cell metadata or content
        
        Args:
            cell: Cell dictionary from notebook
            cell_type: Type of the cell (code, markdown, raw)
            
        Returns:
            str: Comment string or empty string
        """
        comments = []
        
        # Add cell type as comment
        comments.append(f"Cell type: {cell_type}")
        
        # Check for tags in metadata
        metadata = cell.get('metadata', {})
        tags = metadata.get('tags', [])
        if tags:
            comments.append(f"Tags: {', '.join(tags)}")
        
        # Check for execution count for code cells
        if cell_type == 'code':
            execution_count = cell.get('execution_count')
            if execution_count is not None:
                comments.append(f"Execution count: {execution_count}")
        
        # Check for outputs in code cells
        if cell_type == 'code' and cell.get('outputs'):
            output_types = [output.get('output_type', 'unknown') for output in cell.get('outputs', [])]
            if output_types:
                comments.append(f"Output types: {', '.join(set(output_types))}")
        
        return '; '.join(comments) if comments else ""
    
    def to_markdown_table(self, output_path: str = None, cell_type_filter: str = None) -> str:
        """
        Convert the parsed notebook data to a markdown table
        
        Args:
            output_path: Optional path to save the markdown file
            cell_type_filter: Optional filter for cell types ('code', 'markdown', 'raw')
            
        Returns:
            str: Markdown table string
        """
        if not self.cells_data:
            self.parse_notebook(cell_type_filter)
        
        # Create markdown table
        markdown_lines = [
            "# Jupyter Notebook Content",
            "",
            "| Order | Content | Comment |",
            "|-------|---------|---------|"
        ]
        
        for cell in self.cells_data:
            # Escape markdown special characters and handle multiline content
            content = self._escape_markdown(cell.content)
            comment = self._escape_markdown(cell.comment)
            
            # # Truncate very long content for table readability
            # if len(content) > 200:
            #     content = content[:197] + "..."
            
            markdown_lines.append(f"| {cell.order} | {content} | {comment} |")
        
        markdown_content = '\n'.join(markdown_lines)
        
        # Save to file if path provided
        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            print(f"Markdown table saved to: {output_path}")
        
        return markdown_content
    
    def _escape_markdown(self, text: str) -> str:
        """
        Escape markdown special characters and handle multiline text
        
        Args:
            text: Text to escape
            
        Returns:
            str: Escaped text suitable for markdown table
        """
        if not text:
            return ""
        
        # Replace newlines with <br> for table compatibility
        text = text.replace('\n', '<br>')
        
        # Escape pipe characters that would break table formatting
        text = text.replace('|', '\\|')
        
        # Escape other markdown characters
        text = text.replace('*', '\\*')
        text = text.replace('_', '\\_')
        text = text.replace('`', '\\`')
        
        return text
    
    def to_docx(self, output_path: str, cell_type_filter: str = None) -> None:
        """
        Convert the parsed notebook data to a DOCX file
        
        Args:
            output_path: Path to save the DOCX file
            cell_type_filter: Optional filter for cell types ('code', 'markdown', 'raw')
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx library is required for DOCX export. Install with: pip install python-docx")
        
        if not self.cells_data:
            self.parse_notebook(cell_type_filter)
        
        # Create a new document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Jupyter Notebook Content', 0)
        
        # Add table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Set header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Order'
        header_cells[1].text = 'Content'
        header_cells[2].text = 'Comment'
        
        # Add data rows
        for cell_data in self.cells_data:
            row_cells = table.add_row().cells
            row_cells[0].text = str(cell_data.order)
            row_cells[1].text = cell_data.content
            row_cells[2].text = cell_data.comment
        
        # Adjust column widths
        for row in table.rows:
            row.cells[0].width = Inches(0.8)  # Order column
            row.cells[1].width = Inches(4.0)  # Content column
            row.cells[2].width = Inches(2.0)  # Comment column
        
        # Save document
        doc.save(output_path)
        print(f"DOCX file saved to: {output_path}")
    
    def get_summary(self, cell_type_filter: str = None) -> Dict[str, Any]:
        """
        Get a summary of the parsed notebook
        
        Args:
            cell_type_filter: Optional filter for cell types ('code', 'markdown', 'raw')
        
        Returns:
            Dict: Summary statistics
        """
        if not self.cells_data:
            self.parse_notebook(cell_type_filter)
        
        cell_types = {}
        total_chars = 0
        
        for cell in self.cells_data:
            # Extract cell type from comment
            comment_parts = cell.comment.split(';')
            cell_type = "unknown"
            for part in comment_parts:
                if part.strip().startswith("Cell type:"):
                    cell_type = part.split(":")[1].strip()
                    break
            
            cell_types[cell_type] = cell_types.get(cell_type, 0) + 1
            total_chars += len(cell.content)
        
        return {
            "total_cells": len(self.cells_data),
            "cell_types": cell_types,
            "total_characters": total_chars,
            "average_cell_length": total_chars / len(self.cells_data) if self.cells_data else 0
        }
    
    def export_markdown_only(self, markdown_output_path: str = None, docx_output_path: str = None) -> str:
        """
        Convenience method to export only markdown cells
        
        Args:
            markdown_output_path: Optional path to save markdown file
            docx_output_path: Optional path to save DOCX file
            
        Returns:
            str: Markdown table string containing only markdown cells
        """
        # Parse only markdown cells
        self.parse_notebook(cell_type_filter='markdown')
        
        # Generate markdown table
        markdown_content = self.to_markdown_table(markdown_output_path)
        
        # Generate DOCX if path provided and library available
        if docx_output_path and DOCX_AVAILABLE:
            self.to_docx(docx_output_path)
        elif docx_output_path and not DOCX_AVAILABLE:
            print("Warning: python-docx not available, skipping DOCX export")
        
        return markdown_content


def main():
    """
    Main function to demonstrate the notebook parser functionality
    """
    try:
        # Initialize parser
        parser = NotebookParser(nburl)
        
        # Parse the notebook (all cells)
        print("Parsing notebook...")
        cells = parser.parse_notebook()
        print(f"Successfully parsed {len(cells)} cells")
        
        # Print summary
        summary = parser.get_summary()
        print("\nNotebook Summary:")
        print(f"Total cells: {summary['total_cells']}")
        print(f"Cell types: {summary['cell_types']}")
        print(f"Total characters: {summary['total_characters']}")
        print(f"Average cell length: {summary['average_cell_length']:.1f} characters")
        
        # # Generate markdown table (all cells)
        # print("\nGenerating markdown table (all cells)...")
        base_name = os.path.splitext(os.path.basename(nburl))[0]
        markdown_path = f"{base_name}_content.md"
        # markdown_content = parser.to_markdown_table(markdown_path)
        
        # # Generate DOCX file (all cells)
        if DOCX_AVAILABLE:
            print("Generating DOCX file (all cells)...")
            docx_path = f"{base_name}_content.docx"
        #     parser.to_docx(docx_path)
        # else:
        #     print("Skipping DOCX generation (python-docx not available)")
        
        # Generate markdown-only exports
        print("\nGenerating markdown-only exports...")
        markdown_only_path = f"{base_name}_markdown_only.md"
        docx_only_path = f"{base_name}_markdown_only.docx" if DOCX_AVAILABLE else None
        
        # Create a new parser instance for markdown-only export
        parser_md = NotebookParser(nburl)
        markdown_only_content = parser_md.export_markdown_only(
            markdown_output_path=markdown_only_path,
            docx_output_path=docx_only_path
        )
        
        # Get summary for markdown-only cells
        md_summary = parser_md.get_summary()
        print(f"Markdown-only export: {md_summary['total_cells']} cells")
        
        print("\nFirst few cells preview (all cells):")
        for i, cell in enumerate(cells[:3]):
            print(f"\nCell {cell.order}:")
            print(f"Content: {cell.content[:100]}{'...' if len(cell.content) > 100 else ''}")
            print(f"Comment: {cell.comment}")
        
        print(f"\nFiles generated:")
        print(f"- All cells: {markdown_path}, {docx_path if DOCX_AVAILABLE else 'DOCX skipped'}")
        print(f"- Markdown only: {markdown_only_path}, {docx_only_path if DOCX_AVAILABLE else 'DOCX skipped'}")
        
    except Exception as e:
        print(f"Error: {e}")


if __name__ == "__main__":
    main()
